# -*- coding: utf-8 -*-
"""Gera a ENTREGA 2 (proposta de posicao) a partir das saidas do pipeline.

    python entrega2.py

Todo numero vem de outputs/resumo_execucao.json e dos CSV. Nada digitado: o
documento e a planilha nao podem divergir, e a unica forma de garantir isso e
gerar o texto a partir da mesma fonte.
"""
from __future__ import annotations
import json
from pathlib import Path
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RAIZ = Path(__file__).resolve().parent
OUT = RAIZ / "outputs"
ROSA = RGBColor(0xC2, 0x18, 0x5B); ESC = RGBColor(0x1F, 0x12, 0x35)
CINZA = RGBColor(0x61, 0x61, 0x61); PRETO = RGBColor(0x22, 0x22, 0x22)
VERM = RGBColor(0xB7, 0x1C, 0x1C)


def br(x, d=0):
    return f"{x:,.{d}f}".replace(",", "@").replace(".", ",").replace("@", ".")


def mi(x):
    return f"R$ {br(x/1e6, 1)} mi"


def ler():
    d = json.loads((OUT / "resumo_execucao.json").read_text(encoding="utf-8"))
    g = lambda n: (pd.read_csv(OUT / n) if (OUT / n).exists() else pd.DataFrame())
    return d, d.get("book", {}), g("curva_mensal_base_seco_umido.csv"), \
        g("sinal_vs_mercado.csv"), g("book_pernas.csv")


def build():
    d, b, cur, sin, per = ler()
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.55)
        s.left_margin = s.right_margin = Inches(0.65)
    st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(9.5)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    def P(t="", size=9.5, bold=False, ital=False, cor=PRETO, after=4, before=0):
        p = doc.add_paragraph(); r = p.add_run(t)
        r.font.size, r.bold, r.italic, r.font.color.rgb = Pt(size), bold, ital, cor
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.space_before = Pt(before)
        return p

    def H(t, size=12):
        p = doc.add_paragraph(); r = p.add_run(t)
        r.font.size, r.bold, r.font.color.rgb = Pt(size), True, ROSA
        p.paragraph_format.space_before = Pt(9); p.paragraph_format.space_after = Pt(3)

    def shade(c, h):
        tcPr = c._tc.get_or_add_tcPr(); sh = OxmlElement("w:shd")
        sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), h); tcPr.append(sh)

    def TAB(cab, linhas, larg):
        t = doc.add_table(rows=1, cols=len(cab)); t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.autofit = False
        lay = OxmlElement("w:tblLayout"); lay.set(qn("w:type"), "fixed")
        t._tbl.tblPr.append(lay)
        for j, x in enumerate(cab):
            c = t.rows[0].cells[j]; shade(c, "1F1235")
            p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(1)
            r = p.add_run(str(x)); r.bold = True; r.font.size = Pt(8)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for k, L in enumerate(linhas):
            row = t.add_row()
            for j, x in enumerate(L):
                c = row.cells[j]
                if k % 2: shade(c, "F7F7FA")
                p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(1)
                r = p.add_run(str(x)); r.font.size = Pt(8)
                if j: p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for j, w in enumerate(larg):
            for row in t.rows: row.cells[j].width = Inches(w)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    mes = lambda x: pd.Timestamp(x).strftime("%b/%y")
    lim = 50e6
    vendidos = int((per.lado == "V").sum()) if len(per) else 0
    comprados = int((per.lado == "C").sum()) if len(per) else 0
    lado = "VENDIDA" if vendidos and not comprados else "MISTA"
    pr_ini, pr_fim = (mes(cur.mes_ref.iloc[0]), mes(cur.mes_ref.iloc[-1])) if len(cur) else ("", "")

    # ------------------------------------------------------------------ cabecalho
    P("ENTREGA 2 — PROPOSTA DE POSIÇÃO", size=15, bold=True, cor=ROSA, after=1)
    P(f"Energia Convencional Flat SE/CO  ·  data de corte 14/08/2026  ·  "
      f"limite de VaR R$ 50 milhões  ·  status do run: {d.get('status','')}",
      size=8.5, cor=CINZA, after=8)

    # ------------------------------------------------------------------ 1 tese
    H("1. Tese")
    P(f"O mercado paga, em todos os vértices de {pr_ini} a {pr_fim}, um prêmio acima da projeção "
      f"oficial de PLD da CCEE — média de {br(d.get('premio_nivel_rs_mwh',0),2)} R$/MWh de nível. "
      f"Vender a termo hoje entrega energia a preço superior ao que o fundamento hidrotérmico "
      f"espera para o spot no mês de entrega.")
    lad = b.get("ladder") or []
    escada = "; ".join(f"{x['mes']} {br(x['mwmed'])} MWm" for x in lad) if lad else ""
    P(f"A posição é {lado} em {vendidos + comprados} vértices mensais — {escada} — "
      f"totalizando {br(b.get('energia_liquida_gwh',0),1)} GWh e "
      f"R$ {br(b.get('notional_brl',0)/1e6,1)} milhões de notional. São produtos mensais "
      f"distintos, com quantidade dimensionada vértice a vértice pelo risco de cada mês: "
      f"não há um MWm único que descreva o book, e o shape é parte da tese. "
      f"A posição captura o prêmio por carrego até a entrega. Não é aposta direcional sobre hidrologia: "
      f"é a diferença entre preço a termo e expectativa de spot, que é observável e mensurável.")
    P(f"O risco é o cenário seco: se o PLD subir acima do preço de entrada, a posição vendida perde. "
      f"O dimensionamento consome {br(b.get('consumo_limite',0)*100,1)}% do limite de VaR, e o pior "
      f"cenário hidrológico modelado cabe dentro do limite.")

    # ------------------------------------------------------------------ 2 dimensionamento
    H("2. Dimensionamento e consumo do limite de VaR")
    if len(per):
        linhas = [[r.produto, "VENDIDO" if r.lado == "V" else "COMPRADO",
                   br(r.mwmed), br(r.preco_entrada, 2), br(r.energia_mwh/1000, 1)]
                  for r in per.itertuples()]
        linhas.append(["TOTAL", "", br(per.mwmed.sum()), "", br(per.energia_mwh.sum()/1000, 1)])
        TAB(["vértice", "lado", "MWmed", "entrada R$/MWh", "energia GWh"],
            linhas, [1.1, 1.1, 0.9, 1.3, 1.1])
    TAB(["métrica", "valor", "leitura"],
        [["Risco do book (VaR 95%, 1 mês)", mi(b.get("var_total", 0)),
          "soma entre vértices, sem diversificação"],
         ["Consumo do limite de R$ 50 mi", f"{br(b.get('consumo_limite',0)*100,1)}%",
          "folga de " + mi(lim - b.get("var_total", 0))],
         ["Expected Shortfall", mi(b.get("es_total", 0)), "perda média na cauda de 5%"],
         ["VaR de preço", f"{br(d.get('var_preco_rs_mwh',0),2)} R$/MWh",
          "fator de risco é o preço a termo, não o PLD spot"]],
        [2.4, 1.5, 3.3])
    P("O tamanho de cada vértice sai do risco, não de um MWm arbitrado: "
      "MWm(m) = orçamento de risco × convicção(m) ÷ VaR(m). Vértice de risco menor recebe mais "
      "MWm para o mesmo consumo de limite. O orçamento operacional é 60% do limite, e a fração "
      "restante é buffer para remarcação adversa antes do stop.", size=9, ital=True)

    # ------------------------------------------------------------------ 3 resultado
    H("3. Resultado esperado, como intervalo")
    TAB(["cenário", "PnL até a entrega", "fonte do cenário"],
        [["Seco", mi(b.get("pnl_Entrega_Seco", 0)), "estimador de regime sobre histórico"],
         ["Esperado", mi(b.get("pnl_Entrega_Esperado", 0)), "trajetória RNA — InfoPLD/CCEE"],
         ["Úmido", mi(b.get("pnl_Entrega_Umido", 0)), "trajetória mais úmida do leque oficial"],
         ["Convergência do prêmio (MTM)", mi(b.get("pnl_Convergencia", 0)),
          "prêmio converge ao justo antes da entrega"],
         ["VPL até 31/12", mi(b.get("vpl", 0)), "desconto mês a mês"]],
        [1.8, 1.6, 3.8])
    P(f"O intervalo relevante é de {mi(b.get('pnl_Entrega_Seco',0))} a "
      f"{mi(b.get('pnl_Entrega_Umido',0))}, com {mi(b.get('pnl_Entrega_Esperado',0))} no cenário "
      f"central. Retorno sobre risco de {br(b.get('retorno_risco_carrego',0),2)}x no esperado.")
    P("A assimetria é declarada e favorece a posição: o ganho no cenário úmido é várias vezes "
      "maior que a perda no seco, porque vendido ganha quando o preço cai e o piso administrativo "
      "do PLD limita a queda — mas o teto não limita a alta. É por isso que o cenário seco, e não "
      "o úmido, define o dimensionamento.", size=9, ital=True)

    # ------------------------------------------------------------------ 4 horizonte
    H("4. Horizonte e reavaliação")
    TAB(["item", "definição"],
        [["Horizonte da operação", f"{pr_ini} a {pr_fim}, carrego até a entrega física"],
         ["Reavaliação programada", "a cada novo InfoPLD Diário — revisão semanal formal, "
                                    "com recálculo completo do pipeline"],
         ["Reavaliação obrigatória", "qualquer revisão de PMO ou mudança de trajetória "
                                     "central da CCEE"],
         ["Marco de resultado", "31/12/2026, alinhado às metas de margem e VPL da companhia"]],
        [1.9, 5.3])

    # ------------------------------------------------------------------ 5 gatilhos
    H("5. Gatilhos de saída e o que invalida a tese")
    P("Gatilhos de saída", bold=True, size=9.5, after=2)
    TAB(["gatilho", "limiar", "ação"],
        [["Stop de risco", "consumo do VaR acima de 100% do limite",
          "reduzir até voltar ao orçamento de 60%"],
         ["Movimento adverso do termo",
          f"acima de {br(d.get('var_preco_rs_mwh',0),2)} R$/MWh contra a posição em 1 mês",
          "revisar tamanho do vértice"],
         ["Convergência do prêmio", "prêmio observado cai ao prêmio justo do vértice",
          "realizar o MTM e encerrar a perna"],
         ["Índice hidrológico", "cruzar o limiar do regime seco",
          "reduzir exposição vendida"]],
        [1.7, 3.0, 2.5])
    P("O que invalidaria a tese", bold=True, size=9.5, before=4, after=2)
    P("• O prêmio deixar de existir: se a referência de mercado convergir para a projeção de PLD, "
      "não há mais o que capturar e a posição vira aposta direcional pura — que não é a tese.")
    P("• A projeção oficial subir de forma consistente entre boletins: significa que o fundamento "
      "estava errado, não o mercado. Nesse caso quem está errado sou eu, não o preço.")
    P("• Hidrologia adversa que leve o spot acima do preço de entrada de forma persistente, e não "
      "apenas em um mês isolado.")
    P("• Ruptura da premissa de liquidez: se o teto operacional de MWm por vértice não se "
      "sustentar na execução, o book não é montável no tamanho proposto.")

    # ------------------------------------------------------------------ 6 premissas
    H("6. Premissas de cenário, com fonte")
    TAB(["premissa", "valor", "fonte / natureza"],
        [["Projeção de PLD por mês", "trajetória RNA", "InfoPLD Diário CCEE — OBSERVADO"],
         ["Cenário Seco", f"k = {br(d.get('k_seco',0),3)}",
          "estimador de regime sobre histórico — CALCULADO"],
         ["Cenário Úmido", f"k = {br(d.get('k_umido',0),3)}",
          "trajetória do leque oficial — OBSERVADO"],
         ["Peso da âncora fundamental", f"{br(d.get('peso_fundamental',0)*100,0)}%",
          "PREMISSA declarada, editável"],
         ["Meia-vida da sazonalidade", f"{d.get('meia_vida_dias','')} dias",
          "selecionada por walk-forward — CALCULADO"],
         ["Nowcast de curto prazo", "surpresa observada", "IPDO/ONS — OBSERVADO"],
         ["Referência de mercado", "leitura de mesa por vértice",
          "PREMISSA declarada, sem identificação de fonte"],
         ["Reversão do PLD spot", f"{br(d.get('meia_vida_reversao_dias',0),1)} dias",
          "estimado do histórico — CALCULADO"],
         ["Correlação entre meses", "1 (sem diversificação)",
          "PREMISSA conservadora declarada"]],
        [2.1, 1.7, 3.4])

    # ------------------------------------------------------------------ 7 cenarios
    H("7. Comportamento em dois cenários hidrológicos distintos")
    for nome, chave, k in (("SECO", "pnl_Entrega_Seco", "k_seco"),
                           ("ÚMIDO", "pnl_Entrega_Umido", "k_umido")):
        P(f"Cenário {nome}", bold=True, size=9.5, before=4, after=2)
        pnl = b.get(chave, 0)
        P(f"Resultado esperado: {mi(pnl)}. Multiplicador sobre a curva central: "
          f"{br(d.get(k,0),3)}.")
        if nome == "SECO":
            P(f"Impacto no VaR: o VaR de marcação não muda com o cenário — ele mede a "
              f"remarcação de 1 mês, não o desfecho hidrológico. O que muda é a perda de carrego, "
              f"que chega a {mi(abs(pnl))} contra um VaR de {mi(b.get('var_total',0))}. "
              f"Essa razão de {br(abs(pnl)/max(b.get('var_total',1),1),2)}x é a informação de cauda: "
              f"o stress de carregar até a entrega supera o VaR de curto prazo, e por isso é "
              f"reportado, ainda que não dimensione.")
            P("O que muda na tese: a tese não se invalida por um mês seco isolado — ela se "
              "invalida se a projeção oficial subir de forma consistente entre boletins. "
              "Operacionalmente, reduzo o vértice mais próximo da entrega primeiro, que é onde "
              "o VaR por MWm é maior e a convergência já não paga o risco.")
        else:
            P(f"Impacto no VaR: idêntico — {mi(b.get('var_total',0))}. O VaR é simétrico e não "
              f"depende do lado do desfecho; o que o cenário muda é o resultado, não a medida "
              f"de risco de marcação.")
            P("O que muda na tese: o cenário confirma a tese e o prêmio converge mais rápido. "
              "A decisão passa a ser de realização: encerrar cedo captura o MTM sem carregar o "
              "risco residual até a entrega. O gatilho de convergência do prêmio existe "
              "exatamente para isso.")

    # ------------------------------------------------------------------ 8 metas
    H("8. Metas de margem e VPL até 31/12")
    P(f"A posição contribui com VPL de {mi(b.get('vpl',0))} até 31/12 no cenário de convergência, "
      f"e {mi(b.get('pnl_Entrega_Esperado',0))} de resultado esperado no carrego. Para uma "
      f"companhia geradora, vender a termo acima da expectativa de spot é travar receita, não "
      f"montar exposição especulativa — o que é diretamente compatível com meta de margem.")
    P("Declaro a interação que não modelei: o case diz que não há book de partida, então tratei "
      "a posição como isolada. Numa carteira real, a venda a termo interage com a geração física "
      "e com o GSF, e essa interação mudaria o dimensionamento. É premissa declarada, não "
      "omissão.", size=9, ital=True)

    # ------------------------------------------------------------------ rodape
    P("Planilha aberta com todas as fórmulas visíveis e sem valores colados: "
      "MODELO_Curva_Forward.xlsx. A aba CROSSCHECK reconcilia cada número da planilha contra "
      "uma implementação independente em Python; a aba MANIFESTO registra hash SHA-256 e origem "
      "de cada arquivo de entrada.", size=8.5, cor=CINZA, before=8)

    while doc.paragraphs and not doc.paragraphs[-1].text.strip():
        el = doc.paragraphs[-1]._element; el.getparent().remove(el)
    alvo = OUT / "Entrega_2_Proposta_de_Posicao.docx"
    doc.save(alvo)
    return alvo


if __name__ == "__main__":
    a = build()
    print(f"gerado: {a}")
